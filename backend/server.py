from fastapi import FastAPI, APIRouter, HTTPException, Response, Request, Depends, File, UploadFile
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import httpx
from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContent
import base64

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Models
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    name: str
    picture: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class Session(BaseModel):
    model_config = ConfigDict(extra="ignore")
    session_token: str
    user_id: str
    expires_at: str

class AllergyProfile(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    allergies: List[str]
    dietary_restrictions: Optional[List[str]] = []
    religion_restrictions: Optional[List[str]] = []
    skin_sensitivities: Optional[List[str]] = []
    severity_notes: Optional[str] = ""
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class AllergyProfileCreate(BaseModel):
    allergies: List[str]
    dietary_restrictions: Optional[List[str]] = []
    religion_restrictions: Optional[List[str]] = []
    skin_sensitivities: Optional[List[str]] = []
    severity_notes: Optional[str] = ""

class AnalysisRequest(BaseModel):
    query: str

class AnalysisResult(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    query: str
    analysis_type: str
    result: str
    is_safe: bool
    warnings: List[str]
    alternatives: List[str]
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ImageAnalysisResult(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    product_name: Optional[str] = ""
    ingredients: List[str] = []
    detected_allergens: List[str] = []
    is_safe: bool = False
    safety_rating: int = 0  # 0-100 rating for the user
    warnings: List[str] = []
    alternatives: List[str] = []  # Safe alternative products
    detailed_analysis: str = ""
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class MenuDish(BaseModel):
    name: str
    description: Optional[str] = ""
    is_safe: bool
    allergens: List[str]
    warnings: List[str]
    modifications: List[str]

class MenuAnalysisResult(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    restaurant_name: Optional[str] = ""
    source: str  # 'url' or 'photo'
    source_data: str  # URL or 'uploaded_photo'
    safe_dishes: List[MenuDish]
    unsafe_dishes: List[MenuDish]
    summary: str
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class MenuURLRequest(BaseModel):
    url: str

class RecipeRequest(BaseModel):
    food_item: str
    exclude_recipes: Optional[List[str]] = []

class Recipe(BaseModel):
    name: str
    description: Optional[str] = ""
    prep_time: Optional[str] = ""
    cook_time: Optional[str] = ""
    servings: Optional[str] = ""
    ingredients: List[str]
    instructions: List[str]
    allergen_warnings: List[str]
    safe_for_user: bool

class RecipeFinderResult(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    food_item: str
    recipes: List[Recipe]
    summary: str
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Authentication helper
async def get_current_user(request: Request) -> str:
    session_token = request.cookies.get('session_token')
    if not session_token:
        auth_header = request.headers.get('Authorization')
        if auth_header and auth_header.startswith('Bearer '):
            session_token = auth_header.split(' ')[1]
    
    if not session_token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    session = await db.sessions.find_one({"session_token": session_token})
    if not session:
        raise HTTPException(status_code=401, detail="Invalid session")
    
    expires_at = datetime.fromisoformat(session['expires_at'])
    if expires_at < datetime.now(timezone.utc):
        await db.sessions.delete_one({"session_token": session_token})
        raise HTTPException(status_code=401, detail="Session expired")
    
    return session['user_id']

# Health check endpoint
@api_router.get("/")
async def root():
    return {"status": "ok", "message": "ClarifyAI API"}

# Auth endpoints
@api_router.post("/auth/session")
async def create_session(request: Request, response: Response):
    session_id = request.headers.get('X-Session-ID')
    if not session_id:
        raise HTTPException(status_code=400, detail="Session ID required")
    
    async with httpx.AsyncClient() as client:
        try:
            api_response = await client.get(
                "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data",
                headers={"X-Session-ID": session_id}
            )
            if api_response.status_code != 200:
                raise HTTPException(status_code=401, detail="Invalid session")
            
            session_data = api_response.json()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Auth service error: {str(e)}")
    
    # Create or get user
    user_email = session_data['email']
    existing_user = await db.users.find_one({"email": user_email})
    
    if not existing_user:
        user = User(
            email=session_data['email'],
            name=session_data['name'],
            picture=session_data.get('picture')
        )
        await db.users.insert_one(user.model_dump())
        user_id = user.id
    else:
        user_id = existing_user['id']
    
    # Create session
    session_token = session_data['session_token']
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    
    session = Session(
        session_token=session_token,
        user_id=user_id,
        expires_at=expires_at.isoformat()
    )
    await db.sessions.insert_one(session.model_dump())
    
    # Set cookie
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=7 * 24 * 60 * 60,
        path="/"
    )
    
    # Get user data
    user_data = await db.users.find_one({"id": user_id}, {"_id": 0})
    return {"user": user_data, "session_token": session_token}

@api_router.get("/auth/me")
async def get_me(user_id: str = Depends(get_current_user)):
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user

@api_router.post("/auth/logout")
async def logout(response: Response, user_id: str = Depends(get_current_user), request: Request = None):
    session_token = request.cookies.get('session_token')
    if session_token:
        await db.sessions.delete_one({"session_token": session_token})
    
    response.delete_cookie(key="session_token", path="/")
    return {"message": "Logged out successfully"}

# Allergy Profile endpoints
@api_router.post("/profile/allergy", response_model=AllergyProfile)
async def create_allergy_profile(profile: AllergyProfileCreate, user_id: str = Depends(get_current_user)):
    # Delete existing profile
    await db.allergy_profiles.delete_many({"user_id": user_id})
    
    allergy_profile = AllergyProfile(
        user_id=user_id,
        **profile.model_dump()
    )
    await db.allergy_profiles.insert_one(allergy_profile.model_dump())
    return allergy_profile

@api_router.get("/profile/allergy", response_model=AllergyProfile)
async def get_allergy_profile(user_id: str = Depends(get_current_user)):
    profile = await db.allergy_profiles.find_one({"user_id": user_id}, {"_id": 0})
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    return profile

@api_router.put("/profile/allergy", response_model=AllergyProfile)
async def update_allergy_profile(profile: AllergyProfileCreate, user_id: str = Depends(get_current_user)):
    existing = await db.allergy_profiles.find_one({"user_id": user_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Profile not found")
    
    updated_profile = AllergyProfile(
        id=existing['id'],
        user_id=user_id,
        **profile.model_dump()
    )
    await db.allergy_profiles.update_one(
        {"user_id": user_id},
        {"$set": updated_profile.model_dump()}
    )
    return updated_profile

# AI Analysis endpoint
@api_router.post("/analyze", response_model=AnalysisResult)
async def analyze_item(request: AnalysisRequest, user_id: str = Depends(get_current_user)):
    # Get user's allergy profile
    profile = await db.allergy_profiles.find_one({"user_id": user_id})
    if not profile:
        raise HTTPException(status_code=400, detail="Please set up your allergy profile first")
    
    allergies = profile.get('allergies', [])
    dietary_restrictions = profile.get('dietary_restrictions', [])
    religion_restrictions = profile.get('religion_restrictions', [])
    skin_sensitivities = profile.get('skin_sensitivities', [])
    
    # Check if query is a URL
    is_url = request.query.strip().startswith(('http://', 'https://'))
    product_info = ""
    
    if is_url:
        # Fetch and extract product information from URL
        try:
            from bs4 import BeautifulSoup
            import io
            from PyPDF2 import PdfReader
            from docx import Document
            import openpyxl
            
            async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
                response = await client.get(request.query)
                if response.status_code != 200:
                    raise HTTPException(status_code=400, detail="Unable to fetch product page")
                
                content_type = response.headers.get('content-type', '').lower()
                
                # Try to extract from different file formats
                extracted_text = None
                
                # PDF files
                if 'pdf' in content_type:
                    pdf_reader = PdfReader(io.BytesIO(response.content))
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    extracted_text = text
                
                # Word documents
                elif 'word' in content_type or 'document' in content_type:
                    doc = Document(io.BytesIO(response.content))
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    extracted_text = text
                
                # Plain text or HTML
                else:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    
                    # Remove unwanted elements
                    for element in soup(["script", "style", "nav", "header", "footer", "iframe", "noscript"]):
                        element.decompose()
                    
                    # Try to find product-related content
                    product_sections = soup.find_all(['div', 'section', 'article'], class_=lambda x: x and any(
                        keyword in str(x).lower() for keyword in ['product', 'item', 'detail', 'description', 'ingredient', 'content', 'info']
                    ))
                    
                    if product_sections:
                        for section in product_sections:
                            text = section.get_text()
                            lines = (line.strip() for line in text.splitlines())
                            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                            clean_text = '\n'.join(chunk for chunk in chunks if chunk and len(chunk) > 5)
                            if clean_text:
                                product_info += clean_text + "\n\n"
                    
                    # If no specific sections found, get all visible text
                    if not product_info:
                        text = soup.get_text()
                        lines = (line.strip() for line in text.splitlines())
                        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                        product_info = '\n'.join(chunk for chunk in chunks if chunk and len(chunk) > 5)
                
                if extracted_text:
                    product_info = extracted_text
                
                # Limit content size
                product_info = product_info[:15000]
                
                if not product_info or len(product_info) < 50:
                    raise HTTPException(status_code=400, detail="Could not extract product information from URL")
                
        except httpx.RequestError as e:
            logging.error(f"URL fetch error: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Unable to fetch product page: {str(e)}")
        except Exception as e:
            logging.error(f"URL processing error: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing URL: {str(e)}")
    
    # Create AI prompt
    if is_url:
        system_message = f"""You are an expert allergy assistant. Analyze product information extracted from a website for allergy safety.
    
User's allergies: {', '.join(allergies) if allergies else 'None'}
Dietary restrictions: {', '.join(dietary_restrictions) if dietary_restrictions else 'None'}
Religion restrictions: {', '.join(religion_restrictions) if religion_restrictions else 'None'}
Skin sensitivities: {', '.join(skin_sensitivities) if skin_sensitivities else 'None'}
    
Provide a thorough analysis including:
1. Identify the product name and type
2. Extract all ingredients if available
3. Safety assessment (safe/warning/danger) - Focus on ACTUAL INGREDIENTS only
4. Specific concerns related to user's allergies based on listed ingredients
5. Check against religious dietary laws (Halal, Kosher, Hindu vegetarian, etc.)
6. For perfumes/fragrances: identify common allergen compounds (linalool, limonene, citronellol, geraniol, etc.)
7. Alternative suggestions if unsafe - MUST provide 3-5 specific alternatives when item is unsafe

IMPORTANT: Focus on actual ingredients present in the product. Do NOT emphasize cross-contamination warnings as these are often standard disclaimers."""
        
        user_message = f"""Analyze this product information extracted from the URL: {request.query}

Product Information:
{product_info}
    
Provide your response in this JSON format:
{{
  "is_safe": true/false,
  "summary": "Brief safety summary including product name and type",
  "warnings": ["warning1", "warning2"],
  "alternatives": ["alternative1", "alternative2", "alternative3", "alternative4", "alternative5"],
  "detailed_analysis": "Detailed explanation including ingredients found and safety assessment"
}}

IMPORTANT: If is_safe is false, you MUST provide 3-5 safe alternatives that the user can use instead."""
    else:
        system_message = f"""You are an expert allergy assistant. Analyze products, ingredients, foods, perfumes, and fragrances for allergy safety.
    
User's allergies: {', '.join(allergies) if allergies else 'None'}
Dietary restrictions: {', '.join(dietary_restrictions) if dietary_restrictions else 'None'}
Religion restrictions: {', '.join(religion_restrictions) if religion_restrictions else 'None'}
Skin sensitivities: {', '.join(skin_sensitivities) if skin_sensitivities else 'None'}
    
Provide a thorough analysis including:
1. Safety assessment (safe/warning/danger) - Focus on ACTUAL INGREDIENTS only
2. Specific concerns related to user's allergies based on listed ingredients
3. Check against religious dietary laws (Halal, Kosher, Hindu vegetarian, etc.)
4. For perfumes/fragrances: identify common allergen compounds (linalool, limonene, citronellol, geraniol, etc.)
5. Alternative suggestions if unsafe - MUST provide 3-5 specific alternatives when item is unsafe
6. Emergency advice if needed

IMPORTANT: Focus on actual ingredients present in the item. Do NOT emphasize cross-contamination warnings as these are often standard disclaimers. Only mention cross-contamination if it's a severe allergy and truly critical."""
        
        user_message = f"""Analyze this product: {request.query}
    
Provide your response in this JSON format:
{{
  "is_safe": true/false,
  "summary": "Brief safety summary",
  "warnings": ["warning1", "warning2"],
  "alternatives": ["alternative1", "alternative2", "alternative3", "alternative4", "alternative5"],
  "detailed_analysis": "Detailed explanation"
}}

IMPORTANT: If is_safe is false, you MUST provide 3-5 safe alternatives that the user can use instead. These should be specific product names or ingredients that are safe for their allergies."""
    
    try:
        # Initialize Gemini chat - Using your free Google API key
        chat = LlmChat(
            api_key=os.environ.get('GOOGLE_API_KEY', os.environ['EMERGENT_LLM_KEY']),
            session_id=f"analysis_{user_id}_{uuid.uuid4()}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.0-flash-exp")
        
        message = UserMessage(text=user_message)
        ai_response = await chat.send_message(message)
        
        # Parse AI response
        import json
        # Try to extract JSON from response
        response_text = ai_response
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
        
        try:
            parsed = json.loads(response_text)
        except:
            # Fallback if JSON parsing fails
            parsed = {
                "is_safe": False,
                "summary": response_text[:200],
                "warnings": ["Please review the detailed analysis"],
                "alternatives": [],
                "detailed_analysis": response_text
            }
        
        result = AnalysisResult(
            user_id=user_id,
            query=request.query,
            analysis_type="url" if is_url else "text",
            result=parsed.get('detailed_analysis', response_text),
            is_safe=parsed.get('is_safe', False),
            warnings=parsed.get('warnings', []),
            alternatives=parsed.get('alternatives', [])
        )
        
        # Save to history
        await db.analysis_history.insert_one(result.model_dump())
        
        return result
    
    except Exception as e:
        logging.error(f"Analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

# History endpoint
@api_router.get("/history", response_model=List[AnalysisResult])
async def get_history(user_id: str = Depends(get_current_user)):
    history = await db.analysis_history.find(
        {"user_id": user_id},
        {"_id": 0}
    ).sort("timestamp", -1).limit(20).to_list(20)
    return history

# Clear history endpoint
@api_router.delete("/history")
async def clear_history(user_id: str = Depends(get_current_user)):
    result = await db.analysis_history.delete_many({"user_id": user_id})
    return {"message": f"Cleared {result.deleted_count} history items", "deleted_count": result.deleted_count}

# Image Analysis endpoint - Updated for all product types
@api_router.post("/analyze-image", response_model=ImageAnalysisResult)
async def analyze_image(
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user)
):
    # Get user's allergy profile
    profile = await db.allergy_profiles.find_one({"user_id": user_id})
    if not profile:
        raise HTTPException(status_code=400, detail="Please set up your allergy profile first")
    
    allergies = profile.get('allergies', [])
    dietary_restrictions = profile.get('dietary_restrictions', [])
    religion_restrictions = profile.get('religion_restrictions', [])
    skin_sensitivities = profile.get('skin_sensitivities', [])
    
    try:
        # Read image file
        image_bytes = await file.read()
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        
        # Create AI prompt for all product types
        system_message = f"""You are an expert product label analyzer for ALL types of products including food, skincare, cosmetics, fragrances, and personal care products.

User's allergies: {', '.join(allergies) if allergies else 'None'}
Dietary restrictions: {', '.join(dietary_restrictions) if dietary_restrictions else 'None'}
Religion restrictions: {', '.join(religion_restrictions) if religion_restrictions else 'None'}
Skin sensitivities: {', '.join(skin_sensitivities) if skin_sensitivities else 'None'}

Your task:
1. Identify the product type (food, skincare, cosmetic, perfume, cologne, fragrance, etc.)
2. Read all text from the product label image
3. Extract the product name if visible
4. List all ingredients found
5. Identify potential allergens and irritants based on ACTUAL INGREDIENTS:
   - For food: gluten, nuts, dairy, soy, eggs, fish, shellfish, sesame, etc.
   - For skincare/cosmetics: fragrances, parabens, sulfates, alcohol, essential oils, preservatives, etc.
   - For perfumes/colognes: specific fragrance compounds, alcohol content, allergens like linalool, limonene, citronellol, geraniol, etc.
   - Cross-check with user's specific allergies and skin sensitivities
6. Provide safety assessment based on product type
7. Calculate a personalized safety rating (0-100) where:
   - 100 = Completely safe, no concerns
   - 75-99 = Generally safe, minor considerations
   - 50-74 = Caution advised, some concerning ingredients
   - 25-49 = High risk, multiple allergens present
   - 0-24 = Dangerous, contains user's major allergens

IMPORTANT: Focus on actual ingredients in the product. Do NOT overemphasize cross-contamination warnings or "may contain" statements - these are standard disclaimers and don't indicate the product actually contains the allergen.

Respond in JSON format:
{{
  "product_name": "Product name or empty string",
  "ingredients": ["ingredient1", "ingredient2"],
  "detected_allergens": ["allergen1", "allergen2"],
  "is_safe": true/false,
  "safety_rating": 85,
  "warnings": ["warning1", "warning2"],
  "alternatives": ["safe alternative 1", "safe alternative 2", "safe alternative 3", "safe alternative 4", "safe alternative 5"],
  "detailed_analysis": "Detailed explanation including product type, safety assessment, and rating justification"
}}

CRITICAL: If the product is unsafe (is_safe = false or safety_rating < 75), you MUST provide 3-5 specific safe alternative products in the "alternatives" array. These should be real product names or categories that are safe for the user's allergies."""
        
        user_message = "Analyze this product label image. Identify the product type, extract all ingredients, and identify any allergens or irritants based on the user's profile."
        
        # Initialize Gemini chat with vision using FileContent - Using your free Google API key
        chat = LlmChat(
            api_key=os.environ.get('GOOGLE_API_KEY', os.environ['EMERGENT_LLM_KEY']),
            session_id=f"image_analysis_{user_id}_{uuid.uuid4()}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.0-flash-exp")
        
        # Create FileContent for the image
        file_content = FileContent(
            content_type="image/jpeg",
            file_content_base64=image_base64
        )
        
        message = UserMessage(
            text=user_message,
            file_contents=[file_content]
        )
        ai_response = await chat.send_message(message)
        
        # Parse AI response
        import json
        response_text = ai_response
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
        
        try:
            parsed = json.loads(response_text)
        except:
            parsed = {
                "product_name": "Unknown Product",
                "ingredients": [],
                "detected_allergens": [],
                "is_safe": False,
                "safety_rating": 0,
                "warnings": ["Unable to parse label completely. Please review manually."],
                "alternatives": [],
                "detailed_analysis": response_text
            }
        
        result = ImageAnalysisResult(
            user_id=user_id,
            product_name=parsed.get('product_name', ''),
            ingredients=parsed.get('ingredients', []),
            detected_allergens=parsed.get('detected_allergens', []),
            is_safe=parsed.get('is_safe', False),
            safety_rating=parsed.get('safety_rating', 0),
            warnings=parsed.get('warnings', []),
            alternatives=parsed.get('alternatives', []),
            detailed_analysis=parsed.get('detailed_analysis', '')
        )
        
        # Save to history
        await db.image_analysis_history.insert_one(result.model_dump())
        
        return result
    
    except Exception as e:
        error_msg = str(e)
        logging.error(f"Image analysis error: {error_msg}")
        
        # Check for budget/credit errors
        if "budget" in error_msg.lower() or "credit" in error_msg.lower():
            raise HTTPException(
                status_code=402,
                detail="AI credits exhausted. Please add balance to your Universal Key: Go to Profile → Universal Key → Add Balance"
            )
        
        raise HTTPException(status_code=500, detail=f"Image analysis failed: {error_msg}")

# Image History endpoint
@api_router.get("/image-history", response_model=List[ImageAnalysisResult])
async def get_image_history(user_id: str = Depends(get_current_user)):
    history = await db.image_analysis_history.find(
        {"user_id": user_id},
        {"_id": 0}
    ).sort("timestamp", -1).limit(20).to_list(20)
    return history

# Clear image history endpoint
@api_router.delete("/image-history")
async def clear_image_history(user_id: str = Depends(get_current_user)):
    result = await db.image_analysis_history.delete_many({"user_id": user_id})
    return {"message": f"Cleared {result.deleted_count} image history items", "deleted_count": result.deleted_count}

# Menu URL Analysis endpoint
@api_router.post("/analyze-menu-url", response_model=MenuAnalysisResult)
async def analyze_menu_url(
    request: MenuURLRequest,
    user_id: str = Depends(get_current_user)
):
    # Get user's allergy profile
    profile = await db.allergy_profiles.find_one({"user_id": user_id})
    if not profile:
        raise HTTPException(status_code=400, detail="Please set up your allergy profile first")
    
    allergies = profile.get('allergies', [])
    dietary_restrictions = profile.get('dietary_restrictions', [])
    religion_restrictions = profile.get('religion_restrictions', [])
    
    try:
        from bs4 import BeautifulSoup
        from urllib.parse import urljoin, urlparse
        import io
        from PyPDF2 import PdfReader
        from docx import Document
        import openpyxl
        
        all_menu_content = []
        visited_urls = set()
        base_url = request.url
        
        async def extract_content_from_file(content, content_type, filename=""):
            """Extract text from various file formats"""
            try:
                # PDF files
                if 'pdf' in content_type.lower() or filename.endswith('.pdf'):
                    pdf_reader = PdfReader(io.BytesIO(content))
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text
                
                # Word documents
                elif 'word' in content_type.lower() or filename.endswith(('.doc', '.docx')):
                    doc = Document(io.BytesIO(content))
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text
                
                # Excel files
                elif 'excel' in content_type.lower() or 'spreadsheet' in content_type.lower() or filename.endswith(('.xls', '.xlsx')):
                    wb = openpyxl.load_workbook(io.BytesIO(content))
                    text = ""
                    for sheet in wb.worksheets:
                        for row in sheet.iter_rows(values_only=True):
                            row_text = " | ".join([str(cell) if cell else "" for cell in row])
                            if row_text.strip():
                                text += row_text + "\n"
                    return text
                
                # JSON files
                elif 'json' in content_type.lower() or filename.endswith('.json'):
                    import json
                    data = json.loads(content.decode('utf-8'))
                    return json.dumps(data, indent=2)
                
                # CSV files
                elif 'csv' in content_type.lower() or filename.endswith('.csv'):
                    import csv
                    text = ""
                    csv_file = io.StringIO(content.decode('utf-8'))
                    reader = csv.reader(csv_file)
                    for row in reader:
                        text += " | ".join(row) + "\n"
                    return text
                
                # Plain text
                elif 'text' in content_type.lower() or filename.endswith('.txt'):
                    return content.decode('utf-8')
                
                # HTML (existing logic)
                elif 'html' in content_type.lower():
                    return None  # Will be handled by existing HTML logic
                
                # Try to decode as text for unknown formats
                else:
                    try:
                        return content.decode('utf-8')
                    except:
                        return None
                        
            except Exception as e:
                logging.error(f"Error extracting content from {content_type}: {str(e)}")
                return None
        
        async def fetch_and_extract(url, is_main_page=False):
            if url in visited_urls or len(visited_urls) > 10:
                return
            
            visited_urls.add(url)
            
            try:
                async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
                    response = await client.get(url)
                    if response.status_code != 200:
                        return
                    
                    content_type = response.headers.get('content-type', '').lower()
                    
                    # Try to extract from file formats first
                    extracted_text = await extract_content_from_file(response.content, content_type, url)
                    
                    if extracted_text:
                        all_menu_content.append(extracted_text)
                        return
                    
                    # If not a file format or extraction failed, treat as HTML
                    soup = BeautifulSoup(response.text, 'html.parser')
                    
                    # Remove unwanted elements
                    for element in soup(["script", "style", "nav", "header", "footer", "iframe", "noscript"]):
                        element.decompose()
                    
                    # Extract menu content
                    menu_sections = soup.find_all(['main', 'article', 'section', 'div'], class_=lambda x: x and any(
                        keyword in str(x).lower() for keyword in ['menu', 'food', 'dish', 'item', 'product', 'category']
                    ))
                    
                    if menu_sections:
                        for section in menu_sections:
                            text = section.get_text()
                            lines = (line.strip() for line in text.splitlines())
                            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                            clean_text = '\n'.join(chunk for chunk in chunks if chunk and len(chunk) > 10)
                            if clean_text:
                                all_menu_content.append(clean_text)
                    
                    # If this is the main page, find menu-related links to explore
                    if is_main_page and len(visited_urls) < 10:
                        menu_links = []
                        for link in soup.find_all('a', href=True):
                            href = link.get('href')
                            link_text = link.get_text().lower()
                            
                            # Check if link is menu-related
                            menu_keywords = ['menu', 'food', 'dish', 'category', 'breakfast', 'lunch', 'dinner', 
                                           'drink', 'beverage', 'appetizer', 'entree', 'dessert', 'sandwich', 'salad']
                            
                            if any(keyword in link_text or keyword in href.lower() for keyword in menu_keywords):
                                full_url = urljoin(base_url, href)
                                # Only follow links from the same domain
                                if urlparse(full_url).netloc == urlparse(base_url).netloc:
                                    menu_links.append(full_url)
                        
                        # Follow up to 5 menu links
                        for menu_link in menu_links[:5]:
                            await fetch_and_extract(menu_link, is_main_page=False)
                            
            except Exception as e:
                logging.error(f"Error fetching {url}: {str(e)}")
        
        # Start with the main URL
        await fetch_and_extract(base_url, is_main_page=True)
        
        # Combine all menu content
        combined_menu = '\n\n=== MENU SECTION ===\n\n'.join(all_menu_content)
        
        # Filter out promotional content
        filtered_lines = []
        skip_phrases = ['sign in', 'account', 'rewards', 'join', 'login', 'register', 'newsletter', 'subscribe', 'cart', 'checkout']
        for line in combined_menu.split('\n'):
            if len(line) > 10 and not any(phrase in line.lower() for phrase in skip_phrases):
                filtered_lines.append(line)
        
        menu_content = '\n'.join(filtered_lines)[:20000]  # Increased limit for more content
        
        if not menu_content or len(menu_content) < 100:
            raise HTTPException(status_code=400, detail="Could not extract menu content from the website")
        
        # Create AI prompt
        system_message = f"""You are an expert restaurant menu analyzer. You've explored the entire menu across multiple pages.

User's allergies: {', '.join(allergies) if allergies else 'None'}
Dietary restrictions: {', '.join(dietary_restrictions) if dietary_restrictions else 'None'}
Religion restrictions: {', '.join(religion_restrictions) if religion_restrictions else 'None'}

Your task:
1. Review ALL menu items from all sections (appetizers, entrees, desserts, drinks, etc.)
2. Identify the BEST and SAFEST options for this user across the entire menu
3. Focus on dishes that are completely safe based on listed ingredients
4. For potentially unsafe dishes, suggest modifications to make them safe
5. Provide a curated list of recommended items, not everything on the menu

IMPORTANT: Focus on actual ingredients in menu descriptions. Do NOT be overly concerned with cross-contamination as this is typically unavoidable in restaurant kitchens. Only mention if absolutely critical for severe allergies.

Respond in JSON format:
{{
  "restaurant_name": "Restaurant name or empty",
  "safe_dishes": [
    {{
      "name": "Dish name",
      "description": "Description including what makes it great for this user",
      "is_safe": true,
      "allergens": [],
      "warnings": [],
      "modifications": []
    }}
  ],
  "unsafe_dishes": [
    {{
      "name": "Dish name (only include if popular/notable)",
      "description": "Description",
      "is_safe": false,
      "allergens": ["allergen1"],
      "warnings": ["warning"],
      "modifications": ["Order without X", "Ask for Y on the side"]
    }}
  ],
  "summary": "Summary of the best options found across the entire menu for this user's dietary needs"
}}"""
        
        user_message = f"""I've crawled this restaurant's website and found menu content from multiple pages:

{menu_content}

Based on ALL the menu items above, identify the BEST options for this user. Focus on:
- Completely safe dishes (top priority)
- Popular items that can be modified to be safe
- Best drinks and beverages that fit their needs
- Any standout options across all categories

Provide your analysis in the JSON format specified."""
        
        # Initialize Gemini chat - Using your free Google API key
        chat = LlmChat(
            api_key=os.environ.get('GOOGLE_API_KEY', os.environ['EMERGENT_LLM_KEY']),
            session_id=f"menu_url_{user_id}_{uuid.uuid4()}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.0-flash-exp")
        
        message = UserMessage(text=user_message)
        ai_response = await chat.send_message(message)
        
        # Parse AI response
        import json
        response_text = ai_response
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
        
        try:
            parsed = json.loads(response_text)
        except:
            parsed = {
                "restaurant_name": "",
                "safe_dishes": [],
                "unsafe_dishes": [],
                "summary": "Unable to parse menu. Please try a different URL or upload a photo."
            }
        
        result = MenuAnalysisResult(
            user_id=user_id,
            restaurant_name=parsed.get('restaurant_name', ''),
            source='url',
            source_data=request.url,
            safe_dishes=[MenuDish(**dish) for dish in parsed.get('safe_dishes', [])],
            unsafe_dishes=[MenuDish(**dish) for dish in parsed.get('unsafe_dishes', [])],
            summary=parsed.get('summary', '')
        )
        
        # Save to history
        await db.menu_analysis_history.insert_one(result.model_dump())
        
        return result
    
    except httpx.RequestError as e:
        logging.error(f"Menu URL fetch error: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Unable to fetch menu: {str(e)}")
    except Exception as e:
        logging.error(f"Menu URL analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Menu analysis failed: {str(e)}")

# Menu Photo Analysis endpoint
@api_router.post("/analyze-menu-photo", response_model=MenuAnalysisResult)
async def analyze_menu_photo(
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user)
):
    # Get user's allergy profile
    profile = await db.allergy_profiles.find_one({"user_id": user_id})
    if not profile:
        raise HTTPException(status_code=400, detail="Please set up your allergy profile first")
    
    allergies = profile.get('allergies', [])
    dietary_restrictions = profile.get('dietary_restrictions', [])
    religion_restrictions = profile.get('religion_restrictions', [])
    
    try:
        # Read image file
        image_bytes = await file.read()
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        
        # Create AI prompt
        system_message = f"""You are an expert restaurant menu analyzer. Analyze menu items for allergen safety.

User's allergies: {', '.join(allergies) if allergies else 'None'}
Dietary restrictions: {', '.join(dietary_restrictions) if dietary_restrictions else 'None'}
Religion restrictions: {', '.join(religion_restrictions) if religion_restrictions else 'None'}

Your task:
1. Read all text from the menu photo
2. Extract restaurant name if visible
3. List all menu items with descriptions
4. For each dish, determine if it's safe based on user's allergies
5. Identify potential allergens
6. Suggest modifications for unsafe dishes

Respond in JSON format:
{{
  "restaurant_name": "Restaurant name or empty",
  "safe_dishes": [
    {{
      "name": "Dish name",
      "description": "Description",
      "is_safe": true,
      "allergens": [],
      "warnings": [],
      "modifications": []
    }}
  ],
  "unsafe_dishes": [
    {{
      "name": "Dish name",
      "description": "Description",
      "is_safe": false,
      "allergens": ["allergen1"],
      "warnings": ["Contains nuts"],
      "modifications": ["Order without nuts", "Ask about substitutions"]
    }}
  ],
  "summary": "Overall summary with recommendations"
}}"""
        
        user_message = "Analyze this restaurant menu photo. Extract all menu items and provide allergen safety analysis."
        
        # Initialize Gemini chat - Using your free Google API key
        chat = LlmChat(
            api_key=os.environ.get('GOOGLE_API_KEY', os.environ['EMERGENT_LLM_KEY']),
            session_id=f"menu_photo_{user_id}_{uuid.uuid4()}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.0-flash-exp")
        
        # Create FileContent for the image
        file_content = FileContent(
            content_type="image/jpeg",
            file_content_base64=image_base64
        )
        
        message = UserMessage(
            text=user_message,
            file_contents=[file_content]
        )
        ai_response = await chat.send_message(message)
        
        # Parse AI response
        import json
        response_text = ai_response
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
        
        try:
            parsed = json.loads(response_text)
        except:
            parsed = {
                "restaurant_name": "",
                "safe_dishes": [],
                "unsafe_dishes": [],
                "summary": "Unable to parse menu clearly. Please try with a clearer photo."
            }
        
        result = MenuAnalysisResult(
            user_id=user_id,
            restaurant_name=parsed.get('restaurant_name', ''),
            source='photo',
            source_data='uploaded_photo',
            safe_dishes=[MenuDish(**dish) for dish in parsed.get('safe_dishes', [])],
            unsafe_dishes=[MenuDish(**dish) for dish in parsed.get('unsafe_dishes', [])],
            summary=parsed.get('summary', '')
        )
        
        # Save to history
        await db.menu_analysis_history.insert_one(result.model_dump())
        
        return result
    
    except Exception as e:
        logging.error(f"Menu photo analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Menu analysis failed: {str(e)}")

# Menu History endpoint
@api_router.get("/menu-history", response_model=List[MenuAnalysisResult])
async def get_menu_history(user_id: str = Depends(get_current_user)):
    history = await db.menu_analysis_history.find(
        {"user_id": user_id},
        {"_id": 0}
    ).sort("timestamp", -1).limit(20).to_list(20)
    return history

# Clear menu history endpoint
@api_router.delete("/menu-history")
async def clear_menu_history(user_id: str = Depends(get_current_user)):
    result = await db.menu_analysis_history.delete_many({"user_id": user_id})
    return {"message": f"Cleared {result.deleted_count} menu history items", "deleted_count": result.deleted_count}

# Recipe Finder endpoint
@api_router.post("/recipe-finder", response_model=RecipeFinderResult)
async def find_recipes(
    request: RecipeRequest,
    user_id: str = Depends(get_current_user)
):
    # Get user's allergy profile
    profile = await db.allergy_profiles.find_one({"user_id": user_id})
    if not profile:
        raise HTTPException(status_code=400, detail="Please set up your allergy profile first")
    
    allergies = profile.get('allergies', [])
    dietary_restrictions = profile.get('dietary_restrictions', [])
    religion_restrictions = profile.get('religion_restrictions', [])
    skin_sensitivities = profile.get('skin_sensitivities', [])
    
    try:
        # Create AI prompt for recipe generation
        system_message = f"""You are an expert chef and nutritionist specializing in allergy-safe cooking. Generate safe, delicious recipes.

User's allergies: {', '.join(allergies) if allergies else 'None'}
Dietary restrictions: {', '.join(dietary_restrictions) if dietary_restrictions else 'None'}
Religion restrictions: {', '.join(religion_restrictions) if religion_restrictions else 'None'}
Skin sensitivities (avoid if relevant): {', '.join(skin_sensitivities) if skin_sensitivities else 'None'}

Your task:
1. Generate 2-3 different recipe variations for the requested food item
2. Ensure ALL recipes are 100% safe for the user's allergies and restrictions
3. Avoid ALL allergens completely - no substitutions that contain the same allergen
4. Respect dietary and religious restrictions (Halal, Kosher, Vegetarian, Vegan, etc.)
5. Provide complete, detailed recipes with ingredients and step-by-step instructions
6. Include prep time, cook time, and servings
7. List any potential allergen warnings even if they're avoided in the recipe
8. BE CREATIVE AND DIVERSE - think of different cooking styles, cuisines, and variations

Be creative with substitutions and make recipes that are both safe AND delicious. Generate unique variations each time."""
        
        user_message = f"""Please create allergy-safe recipes for: {request.food_item}

{f"IMPORTANT: Do NOT generate any of these recipes as they were already shown: {', '.join(request.exclude_recipes)}" if request.exclude_recipes else ""}

Provide your response in this JSON format with EXACTLY 3 different recipe variations:
{{
  "recipes": [
    {{
      "name": "Recipe name 1",
      "description": "Brief description",
      "prep_time": "15 minutes",
      "cook_time": "30 minutes",
      "servings": "4 servings",
      "ingredients": ["ingredient 1", "ingredient 2", "..."],
      "instructions": ["Step 1", "Step 2", "..."],
      "allergen_warnings": ["Note about allergens"],
      "safe_for_user": true
    }},
    {{
      "name": "Recipe name 2",
      "description": "Brief description",
      "prep_time": "20 minutes",
      "cook_time": "25 minutes",
      "servings": "4 servings",
      "ingredients": ["ingredient 1", "ingredient 2", "..."],
      "instructions": ["Step 1", "Step 2", "..."],
      "allergen_warnings": ["Note about allergens"],
      "safe_for_user": true
    }},
    {{
      "name": "Recipe name 3",
      "description": "Brief description",
      "prep_time": "10 minutes",
      "cook_time": "35 minutes",
      "servings": "4 servings",
      "ingredients": ["ingredient 1", "ingredient 2", "..."],
      "instructions": ["Step 1", "Step 2", "..."],
      "allergen_warnings": ["Note about allergens"],
      "safe_for_user": true
    }}
  ],
  "summary": "Brief summary explaining how these recipes avoid the user's allergens"
}}

CRITICAL: You MUST provide EXACTLY 3 different recipe variations. All recipes MUST be safe for the user's allergies and restrictions.{f" DO NOT include: {', '.join(request.exclude_recipes)}" if request.exclude_recipes else ""}"""
        
        # Initialize Gemini chat
        chat = LlmChat(
            api_key=os.environ.get('GOOGLE_API_KEY', os.environ['EMERGENT_LLM_KEY']),
            session_id=f"recipe_{user_id}_{uuid.uuid4()}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.0-flash-exp")
        
        message = UserMessage(text=user_message)
        ai_response = await chat.send_message(message)
        
        # Parse AI response
        import json
        response_text = ai_response
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
        
        try:
            parsed = json.loads(response_text)
        except:
            # Fallback if JSON parsing fails
            parsed = {
                "recipes": [],
                "summary": "Failed to parse recipes. Please try again."
            }
        
        result = RecipeFinderResult(
            user_id=user_id,
            food_item=request.food_item,
            recipes=[Recipe(**recipe) for recipe in parsed.get('recipes', [])],
            summary=parsed.get('summary', 'Recipes generated based on your allergy profile.')
        )
        
        # Save to history
        await db.recipe_history.insert_one(result.model_dump())
        
        return result
    
    except Exception as e:
        logging.error(f"Recipe finder error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Recipe generation failed: {str(e)}")

# Recipe History endpoint
@api_router.get("/recipe-history", response_model=List[RecipeFinderResult])
async def get_recipe_history(user_id: str = Depends(get_current_user)):
    history = await db.recipe_history.find(
        {"user_id": user_id},
        {"_id": 0}
    ).sort("timestamp", -1).limit(20).to_list(20)
    return history

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()